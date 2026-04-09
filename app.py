import os
import re
import subprocess
from datetime import datetime
from flask import Flask, render_template, request, send_from_directory, flash, jsonify
from docx import Document
from docx.text.paragraph import Paragraph
from copy import deepcopy

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
STATIC_DIR = os.path.join(BASE_DIR, "static")
DOCX_TPL_DIR = os.path.join(BASE_DIR, "docx_templates")
GENERATED_DIR = os.path.join(BASE_DIR, "generated")
INSTANCE_DIR = os.path.join(BASE_DIR, "instance")
os.makedirs(GENERATED_DIR, exist_ok=True)
os.makedirs(INSTANCE_DIR, exist_ok=True)

app = Flask(
    __name__,
    template_folder=TEMPLATES_DIR,
    static_folder=STATIC_DIR,
    instance_path=INSTANCE_DIR,
    instance_relative_config=True,
)
app.secret_key = "dev-secret-key"


# ----------------- helpers -----------------
def sanitize_filename(name: str) -> str:
    name = re.sub(r"[^\w\s\-\.]", "", name, flags=re.U)
    name = re.sub(r"\s+", "_", name).strip("_")
    return name or "Document"

def pack_repeating(prefix: str, fields: list[str]) -> list[dict]:
    lists = {f: request.form.getlist(f"{prefix}_{f}[]") for f in fields}
    length = max((len(v) for v in lists.values()), default=0)
    items = []
    for i in range(length):
        item = {f: (lists[f][i].strip() if i < len(lists[f]) else "") for f in fields}
        if any(item.values()):
            items.append(item)
    return items

def compose_state_country(state: str, country: str) -> str:
    """US -> state only; otherwise 'state, country' (skip empty parts)."""
    s = (state or "").strip()
    c = (country or "").strip()
    if not s and not c:
        return ""
    if c and c.lower() in {"us", "usa", "united states", "united states of america"}:
        return s
    if s and c:
        return f"{s}, {c}"
    return s or c

def replace_in_runs_preserve(paragraph, mapping: dict):
    runs = paragraph.runs
    if not runs:
        return

    # 0) Normalize replacement values first
    def _norm(s: str) -> str:
        if s is None:
            return ""
        # normalize CRLF and CR to LF
        return str(s).replace("\r\n", "\n").replace("\r", "\n")

    norm_mapping = {k: _norm(v) for k, v in mapping.items()}

    # 1) Build full string and index map: full_char_idx -> (run_idx, offset_in_run)
    run_texts = [r.text or "" for r in runs]
    index_map = []
    full_pieces = []
    for ri, t in enumerate(run_texts):
        full_pieces.append(t)
        index_map.extend((ri, oi) for oi in range(len(t)))
    full_text = "".join(full_pieces)
    if not full_text:
        return

    # 2) Find non-overlapping matches (longest keys first)
    keys = sorted(norm_mapping.keys(), key=len, reverse=True)
    occupied = [False] * len(full_text)
    matches = []  # (start, end, replacement)

    for k in keys:
        if not k:
            continue
        klen = len(k)
        start = 0
        repl = norm_mapping[k]
        while True:
            pos = full_text.find(k, start)
            if pos == -1:
                break
            end = pos + klen
            if any(occupied[pos:end]):
                start = pos + 1
                continue
            matches.append((pos, end, repl))
            for i in range(pos, end):
                occupied[i] = True
            start = end  # continue after this match

    if not matches:
        return

    # 3) Apply matches right-to-left
    for start, end, repl in sorted(matches, key=lambda x: x[0], reverse=True):
        ri, oi = index_map[start]
        rj, oj = index_map[end - 1]

        if ri == rj:
            # single-run replacement
            t = runs[ri].text
            runs[ri].text = t[:oi] + repl + t[oj + 1:]
        else:
            # span multiple runs:
            # first run = prefix + replacement
            # middle runs = cleared
            # last run = suffix
            t_first = runs[ri].text
            t_last = runs[rj].text

            runs[ri].text = t_first[:oi] + repl
            for m in range(ri + 1, rj):
                runs[m].text = ""
            runs[rj].text = t_last[oj + 1:]

def replace_in_doc_preserve(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        replace_in_runs_preserve(p, mapping)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_runs_preserve(p, mapping)


# ------------- experience templating (clone with formatting) -------------
def copy_paragraph_with_replacements(before_paragraph, src_paragraph, mapping: dict):
    """
    Insert a deep-copied clone of `src_paragraph` *before* `before_paragraph`.
    Then apply cross-run replacement with CR normalized.
    Returns the new Paragraph.
    """
    # clone the underlying CT_P (XML element)
    new_ctp = deepcopy(src_paragraph._element)
    # insert into document tree
    before_paragraph._element.addprevious(new_ctp)
    # wrap as a python-docx Paragraph object
    new_p = Paragraph(new_ctp, before_paragraph._parent)

    # apply cross-run replacer (with CR normalization inside)
    replace_in_runs_preserve(new_p, mapping)
    return new_p

def find_first_experience_block(doc: Document):
    """
    Find the first experience block as the sequence of paragraphs
    starting from the first paragraph containing '[Company Name]'
    until just before 'SKILLS, ACTIVITIES & INTERESTS' or end-of-doc.
    Return (start_idx, end_idx_exclusive). None,None if not found.
    """
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        if "[Company Name]" in p.text:
            start = i
            break
    if start is None:
        return None, None
    end = len(paras)
    for j in range(start + 1, len(paras)):
        if "SKILLS, ACTIVITIES & INTERESTS" in paras[j].text:
            end = j
            break
    return start, end

def materialize_experiences(doc: Document, experiences: list[dict]):
    """
    Fill the first experience block with experiences[0];
    For experiences[1:], clone the original block after its end,
    preserving formatting and inserting a blank line between blocks.
    """
    if not experiences:
        return
    start, end = find_first_experience_block(doc)
    if start is None:
        return

    # Capture ORIGINAL block paragraphs to clone from
    block_paras = doc.paragraphs[start:end]

    # Build a function that maps one experience to placeholder mapping
    def exp_map(e):
        title = e.get("title", "")
        group = e.get("group", "")
        title_group = f"{title}, {group}" if group else f"{title}"
        mapping = {
            "[Company Name]": e.get("company", ""),
            "[City]": e.get("city", ""),
            "[State/Country]": compose_state_country(e.get("state",""), e.get("country","")),
            "[Position Title], [Group Name]": title_group,
            "[Start Date]": e.get("start", ""),
            "[End Date]": e.get("end", ""),
            "[Experience Description]": e.get("summary", ""),
        }
        return mapping

    block_paras_bak = deepcopy(block_paras)
    # (1) Fill the FIRST block in place (preserve run formatting)
    first_mapping = exp_map(experiences[0])
    for p in block_paras:
        replace_in_runs_preserve(p, first_mapping)

    # (2) For remaining experiences, clone block with mapping
    insert_before = doc.paragraphs[end]
    for e in experiences[1:]:
        # clone each paragraph in original block (not the modified one)
        for src_p in block_paras_bak:
            copy_paragraph_with_replacements(insert_before, src_p, exp_map(e))

        # insert a blank paragraph between experiences
        insert_before = insert_before.insert_paragraph_before("")
        insert_before.style = block_paras_bak[-1].style

def build_documents_from_form(request):
    first_name = request.form.get("first_name","").strip()
    last_name  = request.form.get("last_name","").strip()
    email      = request.form.get("email","").strip()
    phone      = request.form.get("phone","").strip()
    physical_address = request.form.get("physical_address","").strip()  # NEW explicit field
    full_name  = f"{first_name} {last_name}".strip()

    # Education (no ed_desc; degree_type dropdown; SAT vs non-US grade system)
    education = pack_repeating("ed", [
        "school","city","state","country","degree_type","field","start","end",
        "gpa","sat","grade_system","honors","courses"
    ])
    ed = education[0] if education else {}

    # Experiences (unlimited; reuse first experience block)
    experiences = pack_repeating("e", [
        "company","city","state","country","title","group","start","end","summary"
    ])

    # Footer skills area (ensure Languages collected)
    languages_fluent      = request.form.get("languages","").strip()
    languages_conversational = request.form.get("languages_secondary","").strip()
    technical_skills    = request.form.get("technical_skills","").strip()
    certifications      = request.form.get("certifications","").strip()
    activities          = request.form.get("activities","").strip()
    interests           = request.form.get("interests","").strip()

    # Cover Letter fields (Applicant Brief removed; School/Major removed here)
    recruiter_name       = request.form.get("recruiter_name","").strip()
    recruiter_title      = request.form.get("recruiter_title","").strip()
    recruiter_last_name  = recruiter_name.strip().split()[-1] if recruiter_name else ""
    recruiter_salutation = request.form.get("recruiter_salutation","").strip()
    company_name         = request.form.get("company_name","").strip()
    recruiter_address    = request.form.get("recruiter_address","").strip()

    # First paragraph pieces that now come from Education + a plain Year field
    cl_year        = request.form.get("cl_year","").strip()  # still keep a short 'Year' field
    cl_school_name = ed.get("school","")
    cl_major       = ed.get("field","")

    referral_source    = request.form.get("referral_source","").strip()
    firm_impression    = request.form.get("firm_impression","").strip()  # label changed on UI only
    position_name      = request.form.get("position_name","").strip()
    past_experience    = request.form.get("past_experience","").strip()
    experience_theme   = request.form.get("experience_theme","").strip()
    gained_skills      = request.form.get("gained_skills","").strip()
    other_skills       = request.form.get("other_skills","").strip()
    project            = request.form.get("project","").strip()
    project_result     = request.form.get("project_result","").strip()
    background_summary = request.form.get("background_summary","").strip()
    skill_summary      = request.form.get("skill_summary","").strip()
    firm_track_record  = request.form.get("firm_track_record","").strip()
    signature          = request.form.get("signature","").strip()

    # ---------- CV ----------
    cv_doc = Document(os.path.join(DOCX_TPL_DIR, "CV_template.docx"))

    # experiences first (block cloning with preserved formatting)
    materialize_experiences(cv_doc, experiences)

    # header + education simple replacements (run-preserving)
    cv_map = {
        "[Name]": full_name or "",
        "[Physical Address]": physical_address or "",
        "[Phone Number]": phone or "",
        "[Email Address]": email or "",

        "[University Name]": ed.get("school",""),
        "[City]": ed.get("city",""),
        "[State/Country]": compose_state_country(ed.get("state",""), ed.get("country","")),
        "[Arts/Science]": ed.get("degree_type",""),
        "[Major]": ed.get("field",""),
        "[Graduation Date]": ed.get("end",""),
        "[GPA]": ed.get("gpa",""),
        "[SAT]": ed.get("sat",""),
        "[If you’re outside the US, list grades under your system here instead]": ed.get("grade_system",""),
        "[Honors]": ed.get('honors',''),
        "[Economics / Accounting / Finance classes, anything business-related]": ed.get('courses',''),
        "[Fluent]": languages_fluent,
        "[Conversational]":languages_conversational,
        "[List any programming languages – not MS Office/Excel]":technical_skills,
        "[Any extra courses or programs relevant to finance]":certifications,
        "[Student Clubs, Volunteer Work, Independent Activities]":activities,
        "[Keep this to 1-2 lines and be specific; do not go overboard]":interests,
    }
    replace_in_doc_preserve(cv_doc, cv_map)

    fname = sanitize_filename(first_name or "Firstname")
    lname = sanitize_filename(last_name or "Surname")
    cv_docx_path = os.path.join(GENERATED_DIR, f"{fname}_{lname}_CV.docx")
    cv_pdf_path  = os.path.join(GENERATED_DIR, f"{fname}_{lname}_CV.pdf")
    cv_doc.save(cv_docx_path)

    # ---------- Cover Letter ----------
    cl_doc = Document(os.path.join(DOCX_TPL_DIR, "Cover_letter_template.docx"))
    cl_map = {
        "[Your Name]": full_name or "",
        "[Your Address]": physical_address or "",            # ensure mapped
        "[Your Phone Number]": phone or "",
        "[Your Email Address]": email or "",
        "[Date]": datetime.today().strftime("%d %b %Y"),
        "[Name of Recruiter]": recruiter_name or "",
        "[Title]": recruiter_title or "",
        "[Name of Bank]": company_name or "",
        "[Recruiter’s Address]": recruiter_address or "",    # ensure mapped
        "[Mr. / Ms.]": recruiter_salutation or "Mr.",
        "[Recruiter’s Name]":recruiter_last_name or '',

        # first paragraph (school/major now from Education; Year from cl_year)
        "[Year]": cl_year or "",
        "[School Name]": cl_school_name or "",
        "[Major]": cl_major or "",

        # body
        "[Friend / Contact at Firm / Presentation]": referral_source or "",
        "[Your Culture / Working Environment / Bank-Specific Info.]": firm_impression or "",
        "[Investment Banking Analyst / Associate]": position_name or "",

        "[Completed Internships In… / Worked Full-Time In…]": past_experience or "",
        "[Working on Transactions / Leading Teams and Managing Projects / Performing Quantitative Analysis]": experience_theme or "",
        "[Go Into Anything Relevant to Banking, Such As Analytical / Leadership / Teamwork / Finance / Accounting]": gained_skills or "",
        "[Any Other Relevant Skills]": other_skills or "",
        "[High-Impact Project]": project or "",
        "[Describe Results]": project_result or "",

        "[Summarize Internships / Work Experience]": background_summary or "",
        "[Summarize Skills]": skill_summary or "",
        "[Position Name]": position_name or "",
        "[Transactions / Clients]": firm_track_record or "",
        "[Firm Name]": company_name or "",

        # closing
        "[Phone Number]": phone or "",
        "[Email Address]": email or "",
        "[Your Name]": full_name or "",
        "[Signature]": signature or "",
    }
    replace_in_doc_preserve(cl_doc, cl_map)

    cl_docx_path = os.path.join(GENERATED_DIR, f"{fname}_{lname}_CoverLetter.docx")
    cl_pdf_path  = os.path.join(GENERATED_DIR, f"{fname}_{lname}_CoverLetter.pdf")
    cl_doc.save(cl_docx_path)

    # optional PDF export
    def try_docx2pdf(input_docx: str, output_pdf: str) -> bool:
        try:
            from docx2pdf import convert
            convert(input_docx, output_pdf)
            return os.path.exists(output_pdf)
        except Exception:
            return False

    def try_libreoffice(input_docx: str, output_pdf: str) -> bool:
        try:
            outdir = os.path.dirname(output_pdf)
            os.makedirs(outdir, exist_ok=True)
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, input_docx],
                check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
            )
            return os.path.exists(output_pdf)
        except Exception:
            return False

    cv_ok = try_docx2pdf(cv_docx_path, cv_pdf_path) or try_libreoffice(cv_docx_path, cv_pdf_path)
    cl_ok = try_docx2pdf(cl_docx_path, cl_pdf_path) or try_libreoffice(cl_docx_path, cl_pdf_path)

    return {
        "cv_docx": os.path.basename(cv_docx_path),
        "cl_docx": os.path.basename(cl_docx_path),
        "cv_pdf": os.path.basename(cv_pdf_path) if cv_ok else None,
        "cl_pdf": os.path.basename(cl_pdf_path) if cl_ok else None,
        "pdf_ok": cv_ok and cl_ok
    }

# ----------------- routes -----------------
@app.route("/")
def index():
    return render_template("form.html")

@app.route("/generate", methods=["POST"])
def generate():
    # Header
    result=build_documents_from_form(request)
    if result.get('pdf_ok'):
        flash("PDF export failed (DOCX generated). Install Word (docx2pdf) or LibreOffice (soffice) for PDF.", "warning")

    return render_template(
        "success.html",
        cv_pdf = result.get('cv_pdf'),
        cl_pdf = result.get('cl_pdf'),
        cv_docx = result.get('cv_docx'),
        cl_docx = result.get('cl_docx'),
    )

@app.route("/preview", methods=["POST"])
def preview():
    result = build_documents_from_form(request)
    return render_template(
        "preview.html",
        cv_pdf=result.get("cv_pdf"),
        cl_pdf=result.get("cl_pdf"),
        cv_docx=result.get("cv_docx"),
        cl_docx=result.get("cl_docx"),
        pdf_ok=result.get("pdf_ok", False),
    )

@app.route("/download/<path:filename>")
def download(filename):
    return send_from_directory(GENERATED_DIR, filename, as_attachment=True)

@app.route("/pdfshow/<path:filename>")
def pdfshow(filename):
    return send_from_directory(GENERATED_DIR, filename, as_attachment=False)

@app.route("/health")
def health():
    return {"status": "ok"}

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
